"""
Script written and finalized 9/2/2025
John Zahn
This script is modified to eliminate Chromium dependencies, and streamline the process of
creating a critique report for JMATS Training.  The script is designed to be run by non-technical
users, and provides a GUI for file selection and report generation.  The script processes an Excel report
downloaded from lms.c130j.com, cleans the data, removes duplicates, and generates a comprehensive
report in Word format, complete with charts and tables. When pushed to fully embedded application with pyinstaller 
errors were encountered with Kaleido and Chromium dependencies.  This version eliminates those dependencies
and uses Matplotlib for pie and bar charts. 

"""
"""
Import desired modules
"""
import os
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
import plotly.express as px
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openpyxl
import hashlib
import logging
import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
import threading
import time
"""
Create log for error tracking, management, and analysis
"""
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("pycritique.log"),
        logging.StreamHandler()])

logger = logging.getLogger(__name__)

"""
Reusable Functions
""" 


def get_file(prompt):
    root = tk.Tk()
    root.withdraw()
    return filedialog.askopenfilename(title=prompt)


def critReport(x):  # pulls data from lms.c13oj.com
     # sets current directory to desired folder
    logger.info(f'Crit{os.getcwd()}')
    ext = os.path.splitext(x)[1].lower()
    engine = 'xlrd' if ext == '.xls' else 'openpyxl'
    df = pd.read_excel(x, header=None, engine=engine)
     
    logger.info('Dataframe created')
    logger.info(f'Loaded {len(df)} total records (including header/footer rows)')
    initial_rows = len(df)
    df =  df.drop(df.iloc[[0,1,2,3]].index)  # Drop first 4 rows
    df.columns = df.iloc[0]  # Set headers
    df = df.iloc[1:].reset_index(drop=True)  # Drop the header row from data
    df.columns = df.columns.str.strip().str.replace(' ', '').str.lower()
    cols_to_drop = list(df.columns[[4, 5, 6, 7, 12, 13]]) 
    df = df.drop(columns=cols_to_drop)

    """ 
    Remove duplicate entries based on identity and question.  These columns were selected 
    as primary indicators for duplicates.  
    """
    headcount = df.groupby(['firstname', 'lastname', 'question', 'responsedate']).size()
    total_unique_combinations = headcount.shape[0]
    logger.info(total_unique_combinations)
    unique_count = df.groupby(['firstname', 'lastname'])['question'].nunique()
    logger.info(unique_count)
    logger.info(unique_count.shape[0])
    df = df.sort_values(by='curriculum', ascending=False).drop_duplicates(subset=
                                                                          ['firstname', 
                                                                           'lastname', 
                                                                           'question', 
                                                                           'responsedate']).reset_index(drop=True)
    unique_count = df.groupby(['firstname', 'lastname'])['question'].nunique()
    cleaned_rows = len(df)
    logger.info(unique_count.shape[0])
    logger.info('DataFrame Cleaned for Results')
    logger.info(f'Removed duplicates; {initial_rows - len(df)} rows dropped. Final: {len(df)} rows.')
    return df, initial_rows, cleaned_rows
"""
Function to create tablle and set as header as 2 columns, one for logo one for text
"""
def add_logo_and_title(doc, logo_path, title_text):
    # Create a table with 1 row and 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4.5)

    # Logo in left cell
    cell_logo = table.cell(0, 0)
    paragraph_logo = cell_logo.paragraphs[0]
    run_logo = paragraph_logo.add_run()
    if os.path.exists(logo_path):
        run_logo.add_picture(logo_path, width=Inches(0.5))

    # Title in right cell
    cell_text = table.cell(0, 1)
    paragraph_text = cell_text.paragraphs[0]
    paragraph_text.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_text = paragraph_text.add_run(title_text)
    run_text.bold = True
    run_text.italic = True
    run_text.font.size = doc.styles['Heading 1'].font.size

"""
Create table for Question Comments includes comments, curriculum, and score font change color below 3, and note added
"""

def add_comments_table(doc, df, title):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    col_widths = [Inches(4), Inches(2), Inches(0.8)]

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        if i < len(col_widths):
            hdr_cells[i].width = col_widths[i]

    for row_idx, (_, row) in enumerate(df.iterrows()):
        row_cells = table.add_row().cells
        full_row_bold_red = False

        try:
            score_val = float(row.iloc[2])
            if score_val < 3:
                full_row_bold_red = True
        except (ValueError, IndexError, TypeError):
            pass

        for i, val in enumerate(row):
            text = str(val) if pd.notna(val) else ''
            para = row_cells[i].paragraphs[0]

            if full_row_bold_red:
                if i == 0:
                    # Main comment in red + bold
                    main_run = para.add_run(text)
                    main_run.bold = True
                    main_run.font.color.rgb = RGBColor(255, 0, 0)

                    # Append normal warning note
                    note_run = para.add_run('\n⚠️ Score below 3')
                else:
                    run = para.add_run(text)
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                run = para.add_run(text)

            if i < len(col_widths):
                row_cells[i].width = col_widths[i]

        # Alternate row shading
        if not full_row_bold_red and row_idx % 2 == 0:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F2F2F2')
            for cell in row_cells:
                cell._tc.get_or_add_tcPr().append(shading_elm)

""" 
Creates bar graph for each question, color coded based off score
"""


def crit_bar(scorecard_df, question_title, img_format="png", width=800, height=600, scale=1):
    # Define custom colors for scores 1 through 5
    custom_colors = {
        "1": 'red',
        "2": 'yellow',
        "3": '#b7ff0f',
        "4": '#58f15f',
        "5": '#00d146'
    }

    # Set up figure size in inches (width, height), converted from pixels
    dpi = 100
    fig_width = width / dpi
    fig_height = height / dpi

    fig, ax = plt.subplots(figsize=(fig_width, fig_height), dpi=dpi)

    # Get data
    x = scorecard_df['ResponseText'].astype(str)
    y = scorecard_df['Frequency']

    # Assign colors using the custom map
    bar_colors = [custom_colors.get(val, 'gray') for val in x]

    # Create bar chart
    bars = ax.bar(x, y, color=bar_colors)

    # Add text labels
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height + 0.5, str(int(height)),
                ha='center', va='bottom', fontsize=10)

    # Titles and labels
    ax.set_title(question_title, fontsize=14)
    ax.set_xlabel("Score")
    ax.set_ylabel("Count")
    ax.set_ylim(0, max(y) * 1.2)  # Add some padding on Y-axis

    plt.tight_layout()

    # Save to in-memory buffer
    buf = BytesIO()
    fig.savefig(buf, format=img_format, bbox_inches='tight')
    buf.seek(0)

    return fig, buf


""""
Creates panda series pulling data for each specific question, turns into df for future use.
Used with add_comments_table and crit_bar
"""
def question_table(crit): 
    questions = crit['question'].dropna().astype(str).unique()
    questions = [q for q in questions if q not in [
    'Identify your crew position:', 
    'Overall, this refresher course was:']] # consolidated one line iteration and removal based on condition
    
    
    scorecard_results = {}
    comment_results = {}
    bar_chart_results = {}
    
    for q in questions: 
        question_df = crit[crit['question'] == q]
        all_scores = ["1", "2", "3", "4", "5"]

        # Scorecard
        scorecard_raw = question_df['responsetext'].value_counts()
        scorecard = scorecard_raw.reindex(all_scores, fill_value=0).reset_index()
        scorecard.columns = ['ResponseText', 'Frequency']
        scorecard_results[q] = scorecard

        # Comments
        comments_df = pd.DataFrame({
            'Comments': question_df['responsecomments'],
            'Curriculum': question_df['curriculum'],
            'Score': question_df['responsetext']
        })
        last_row = pd.DataFrame([{'Comments': 'LastEntry', 'Curriculum': None, 'Score': None}])
        comments_df = pd.concat([comments_df, last_row], ignore_index=True)
        comments_df = comments_df[comments_df['Comments'].notna()]
        comments_df = comments_df[comments_df['Comments'].str.strip() != ""]
        # Convert Score to numeric for proper sorting (NaNs are pushed last)
        comments_df['Score'] = pd.to_numeric(comments_df['Score'], errors='coerce')
        comments_df = comments_df.sort_values(by='Score', na_position='last')
        comment_results[q] = comments_df
        

        # Bar Chart
        bar_charts = crit_bar(scorecard, q)
        bar_chart_results[q] = bar_charts

        
    return scorecard_results, comment_results, bar_chart_results

"""
Creates table for studnets with no assigned curriculum, allows drill down to ensure complancce with 
Performace work Statement for critique acccomplishment 
"""
def add_unknown_course_table(doc, df, title='Unknown Course Entries'):
    doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("No unknown course entries found.")
        return

    # Columns to show
    cols = ['firstname', 'lastname', 'responsetext']
    df = df[cols].dropna(how='all')
    
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = col.capitalize()

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            row_cells[i].text = str(row[col]) if pd.notna(row[col]) else ""
"""
Creates safe virtial name, ensuring appropriate size and characters
"""
def safe_filename(text):
    return hashlib.md5(text.encode('utf-8')).hexdigest()
"""
Culminates all functions into one, reusable functtion compiles all charts and figures into 
monthly/quarterly report
"""    
def export_to_word(bar_charts, comment_tables, filename= None,
                   totals=None, tbc=None, no_course=None, overall_comments_df=None,
                   initial_rows=None, cleaned_rows=None, du=None, month = None):
    doc = Document()
        # Add logo at title top
    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path = os.path.join(script_dir, 'Nova.png')
    add_logo_and_title(doc, logo_path, 'Critique Results\nJMATS Training')

    doc.add_paragraph (f'This report is consolidated data from lms.c130j.com. It provides monthly or quarterly'
                       f'insight to interested parties (instructors, leadership, or government) to enhance and'
                       f'capitalize on training opportunities at DyessJMATS.  The data obtained is cleaned, sorted,'
                       f'and compiled in Python Script. During the cleaning process duplicated data was removed, unknown'
                       f'exceptions were cleared, and then centralized as a clean data set.\n'
                       f' '
                       f'Data was downloaded on: {datetime.now().date()}\n'
                       f'Initial Record Size: {initial_rows}\n'
                       f'Cleaned Record Size: {cleaned_rows}\n'
                       f'Questions or clarifications may be referred to Site Lead/Training')
   # ===== First Page: Student Overview =====
    if totals is not None and no_course is not None:
        pie_image = 'student_summary.png'
        data = totals.T
        data.columns = ['Count']

        labels = data.index.tolist()
        counts = data['Count'].tolist()
        colors = plt.cm.Set3.colors[:len(counts)]  # Limit colors to number of slices

        fig, ax = plt.subplots(figsize=(6, 6))
        wedges, texts, autotexts = ax.pie(counts,
                                          labels=None,
                                          autopct='%1.1f%%',
                                          startangle=90,
                                          colors=colors)

        legend_labels = [f"{label} ({count})" for label, count in zip(labels, counts)]
        ax.legend(wedges, legend_labels, title="Categories",
                  loc="center left", bbox_to_anchor=(1, 0.5))

        ax.set_title('Critique Summary - Student Overview', fontsize=14)
        plt.tight_layout()

        # Save figure to an in-memory buffer
        buffer = BytesIO()
        fig.savefig(buffer, format='png', bbox_inches='tight')
        buffer.seek(0)

        # Insert directly into your Word document
        doc.add_heading('Student Overview', level=1)
        doc.add_paragraph('This data is pulled directly from LMS.')
        doc.add_picture(buffer, width=Inches(6.5))

        # Cleanup
        buffer.close()
        plt.close(fig)




    # ===== Course Overview =====


    if tbc is not None and no_course is not None:
        data2 = tbc.T
        data2.columns = ['Count']

        fig, ax = plt.subplots(figsize=(10, 6))
        data2.plot(kind='bar', legend=False, ax=ax, color='darkgreen')
        ax.set_title('Critique Summary - Course Overview', fontsize=14)
        ax.set_ylabel('Number of Students')
        ax.set_xlabel('')

        plt.xticks(rotation=45, ha='right')
        for i, count in enumerate(data2['Count']):
            ax.text(i, count + 0.5, str(count), ha='center', va='bottom', fontsize=9)

        plt.tight_layout()

        # Save to buffer instead of disk
        buffer = BytesIO()
        fig.savefig(buffer, format='png', bbox_inches='tight')
        buffer.seek(0)
        plt.close(fig)

        doc.add_heading('Course Overview', level=1)
        doc.add_picture(buffer, width=Inches(6.5))

        buffer.close()

    # Add the unknown course table here
    add_unknown_course_table(doc, no_course)
    doc.add_page_break()
        

    # ===== Overall Comments Table =====
    if overall_comments_df is not None:
        overall = overall_comments_df.copy()
        overall = overall.dropna(how='all')
        overall = overall[~(overall == "").all(axis=1)]
        add_comments_table(doc, overall, 'Overall, this refresher course was:')

    # ===== Bar Charts and Comment Tables by Question =====
    for question in bar_charts.keys():
        doc.add_page_break()
        question_title = str(question) if pd.notna(question) else 'Unknown Question'
        doc.add_heading(question_title, level=1)

        fig, buf = bar_charts[question]  # from your updated crit_bar   
        doc.add_picture(buf, width=Inches(6))
        buf.close()
        

        # Add comments
        if question in comment_tables:
            df = comment_tables[question].copy()
            df = df.dropna(how='all')
            df = df[~(df == '').all(axis=1)]
            add_comments_table(doc, df, 'Comments')

    # ===== Save Document =====
    doc.save(filename)
    logger.info(f'Exported to {filename}')

"""
Executes file as desirred, CLI for user choice of files, and month selection. Ran by .bat file
"""
def run_gui():
    import tkinter.messagebox as messagebox

    root = tk.Tk()
    root.title("Critique Report Generator")
    root.geometry("800x500")

    critique_file_var = tk.StringVar()
    filename_var = tk.StringVar()

    # Progress Bar
    pb = ttk.Progressbar(root, orient='horizontal', length=400, mode='determinate', maximum=100)
    pb.pack(pady=10)

    def browse_critique():
        filename = filedialog.askopenfilename(
            title="Select the Critique excel (.xls, .xlsx) file",
            filetypes=[("Excel files", "*.xls *.xlsx")]
        )
        critique_file_var.set(filename)

    def browse_save_location():
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Save Report As"
        )
        filename_var.set(file_path)

    def generate_report():
        pb['value'] = 0
        root.update_idletasks()

        def task():
            try:
                crit_file = critique_file_var.get()
                output_file = filename_var.get()

                if not crit_file:
                    messagebox.showerror("Missing File", "Please select a critique file.")
                    return

                if not output_file:
                    messagebox.showerror("Missing Output Location", "Please select an output file location.")
                    return

                pb['value'] = 10
                root.update_idletasks()

                # Process Critique
                crit, initial_rows, cleaned_rows = critReport(crit_file)
                pb['value'] = 40
                root.update_idletasks()

                # Aggregations (same as in your script)
                students = crit[['firstname', 'lastname', 'responsedate']].drop_duplicates().shape[0]
                p = crit[crit['responsetext'] == 'Pilot'].shape[0]
                l = crit[(crit['responsetext'] == 'Loadmaster') & (~crit['curriculum'].str.contains('engine ground', case=False, na=False))].shape[0]
                mx = crit[crit['curriculum'].str.contains('engine ground', case=False, na=False) &
                          crit['question'].str.contains('knowledgeable', case=False, na=False)].shape[0]

                PDC = crit[crit['curriculum'].str.contains('pilot block difference.*block 8.1', case=False, na=False) &
                           crit['question'].str.contains('knowledgeable', case=False, na=False)].shape[0]

                LDC = crit[crit['curriculum'].str.contains('loadmaster.*block 8.1', case=False, na=False) &
                           crit['question'].str.contains('knowledgeable', case=False, na=False)].shape[0]

                PSR = crit[crit['curriculum'].str.contains('C-130J Pilot Refresher', case=False, na=False) &
                           crit['question'].str.contains('knowledgeable', case=False, na=False)].shape[0]

                LRT = crit[crit['curriculum'].str.contains('loadmaster refresher', case=False, na=False) &
                           crit['question'].str.contains('knowledgeable', case=False, na=False)].shape[0]

                MX = crit[crit['curriculum'].str.contains('engine ground', case=False, na=False) &
                          crit['question'].str.contains('knowledgeable', case=False, na=False)].shape[0]

                UNK = crit[crit['curriculum'].isna() & crit['question'].str.contains('overall', case=False, na=False)].shape[0]

                NoCourse = crit[crit['curriculum'].isna() & crit['question'].str.contains('identify', case=False, na=False)]

                Totals = pd.DataFrame([{'Pilots': p, 'Loadmasters': l, 'MX': mx}])
                Totals_by_Course = pd.DataFrame([{
                    'PDC': PDC, 'LDC': LDC, 'PSR': PSR, 'LRT': LRT, 'MX': MX, 'Unknown': UNK
                }])

                overall_comments = crit[crit['question'] == "Overall, this refresher course was:"]

                pb['value'] = 70
                root.update_idletasks()

                scorecards, comments, bars = question_table(crit)

                export_to_word(
                    bar_charts=bars,
                    comment_tables=comments,
                    filename=output_file,
                    totals=Totals,
                    tbc=Totals_by_Course,
                    no_course=NoCourse,
                    overall_comments_df=pd.DataFrame({
                        'Comments': overall_comments['responsetext'],
                        'Curriculum': overall_comments['curriculum'],
                    }),
                    initial_rows=initial_rows,
                    cleaned_rows=cleaned_rows
                )

                pb['value'] = 100
                root.update_idletasks()

                messagebox.showinfo("Success", f"Report successfully generated as:\n{output_file}")

            except Exception as e:
                logger.error(f"❌ Error: {e}")
                messagebox.showerror("Error", f"An error occurred:\n{e}")

        threading.Thread(target=task).start()

    # Layout GUI
    tk.Label(root, text="Critique File (required):").pack(pady=5)
    tk.Entry(root, textvariable=critique_file_var, width=60).pack()
    tk.Button(root, text="Browse", command=browse_critique).pack(pady=5)

    tk.Label(root, text="Output File Location:").pack(pady=5)
    tk.Entry(root, textvariable=filename_var, width=60).pack()
    tk.Button(root, text="Browse", command=browse_save_location).pack(pady=5)

    tk.Button(root, text="Generate Report", command=generate_report).pack(pady=20)

    root.mainloop()



if __name__ == "__main__":
    run_gui()

